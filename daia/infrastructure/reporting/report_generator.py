"""
DAIA Report Generator - Reportes profesionales ISO-friendly

Genera reportes en PDF y DOCX con estructura empresarial:
1. Resumen Ejecutivo
2. Hallazgos Críticos
3. Métricas Clave
4. Recomendaciones
5. Conclusión Operativa
"""

import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
from dataclasses import dataclass

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx no disponible. Reportes DOCX deshabilitados.")

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("reportlab no disponible. Reportes PDF deshabilitados.")

from daia.application.services import BatchAuditResult
from daia.domain.models import AuditResult, Finding, FindingSeverity

logger = logging.getLogger(__name__)


@dataclass
class ReportConfig:
    """Configuración de reportes"""
    company_name: str = "DAIA 2.0 - Call Audit System"
    logo_path: Optional[str] = None
    output_dir: str = "reports"
    include_transcripts: bool = False
    language: str = "es"


class ReportGenerator:
    """
    Generador de reportes profesionales
    
    Estructura ISO-friendly para venta empresarial:
    - Resumen Ejecutivo (decisores)
    - Hallazgos Críticos (acción inmediata)
    - Métricas Clave (KPIs medibles)
    - Recomendaciones (valor agregado)
    - Conclusión Operativa (siguiente paso)
    """
    
    def __init__(self, config: Optional[ReportConfig] = None):
        """
        Inicializa el generador
        
        Args:
            config: Configuración de reportes
        """
        self.config = config or ReportConfig()
        self.output_dir = Path(self.config.output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"✓ ReportGenerator inicializado (PDF={PDF_AVAILABLE}, DOCX={DOCX_AVAILABLE})")
    
    def generate_batch_report(
        self,
        batch_result: BatchAuditResult,
        format: str = "both"
    ) -> Dict[str, str]:
        """
        Genera reporte consolidado de batch
        
        Args:
            batch_result: Resultado del batch audit
            format: 'pdf', 'docx', o 'both'
            
        Returns:
            Dict con paths de los reportes generados
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"batch_audit_report_{timestamp}"
        
        generated = {}
        
        if format in ['pdf', 'both'] and PDF_AVAILABLE:
            pdf_path = self.output_dir / f"{base_filename}.pdf"
            self._generate_batch_pdf(batch_result, pdf_path)
            generated['pdf'] = str(pdf_path)
            logger.info(f"✓ PDF generado: {pdf_path.name}")
        
        if format in ['docx', 'both'] and DOCX_AVAILABLE:
            docx_path = self.output_dir / f"{base_filename}.docx"
            self._generate_batch_docx(batch_result, docx_path)
            generated['docx'] = str(docx_path)
            logger.info(f"✓ DOCX generado: {docx_path.name}")
        
        if not generated:
            logger.warning("⚠️ No se generaron reportes (dependencias faltantes)")
        
        return generated
    
    def generate_individual_report(
        self,
        audit_result: AuditResult,
        format: str = "both"
    ) -> Dict[str, str]:
        """
        Genera reporte individual de un audit
        
        Args:
            audit_result: Resultado individual
            format: 'pdf', 'docx', o 'both'
            
        Returns:
            Dict con paths de los reportes generados
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        call_id = str(audit_result.audited_call.call_id).replace('/', '_').replace('\\', '_')
        base_filename = f"audit_{call_id}_{timestamp}"
        
        generated = {}
        
        if format in ['pdf', 'both'] and PDF_AVAILABLE:
            pdf_path = self.output_dir / f"{base_filename}.pdf"
            self._generate_individual_pdf(audit_result, pdf_path)
            generated['pdf'] = str(pdf_path)
            logger.info(f"✓ PDF individual: {pdf_path.name}")
        
        if format in ['docx', 'both'] and DOCX_AVAILABLE:
            docx_path = self.output_dir / f"{base_filename}.docx"
            self._generate_individual_docx(audit_result, docx_path)
            generated['docx'] = str(docx_path)
            logger.info(f"✓ DOCX individual: {docx_path.name}")
        
        return generated
    
    # === PDF GENERATION ===
    
    def _generate_batch_pdf(self, batch_result: BatchAuditResult, output_path: Path):
        """Genera PDF del batch consolidado"""
        if not PDF_AVAILABLE:
            raise RuntimeError("reportlab no disponible")
        
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch
        )
        
        story = []
        styles = getSampleStyleSheet()
        
        # Estilos personalizados
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#2C3E50'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#34495E'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # 1. PORTADA
        story.append(Spacer(1, 2*inch))
        story.append(Paragraph(self.config.company_name, title_style))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph("REPORTE CONSOLIDADO DE AUDITORÍA", styles['Heading2']))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(
            f"Generado: {batch_result.generated_at.strftime('%d/%m/%Y %H:%M')}",
            styles['Normal']
        ))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(
            f"Total de llamadas auditadas: {batch_result.total_calls}",
            styles['Normal']
        ))
        story.append(PageBreak())
        
        # 2. RESUMEN EJECUTIVO
        story.append(Paragraph("1. RESUMEN EJECUTIVO", heading_style))
        story.append(Spacer(1, 0.2*inch))
        
        summary_data = [
            ['Métrica', 'Valor'],
            ['Total de Llamadas', str(batch_result.total_calls)],
            ['Llamadas Aprobadas', str(batch_result.passed_calls)],
            ['Llamadas Rechazadas', str(batch_result.failed_calls)],
            ['Tasa de Aprobación', f"{batch_result.approval_rate:.1f}%"],
            ['QA Score Promedio', f"{batch_result.avg_qa_score:.1f}%"],
            ['Hallazgos Críticos', str(batch_result.critical_findings_count)],
            ['Tiempo de Procesamiento', f"{batch_result.processing_time_seconds:.1f}s"],
        ]
        
        summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498DB')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(summary_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Interpretación
        interpretation = self._get_batch_interpretation(batch_result)
        story.append(Paragraph("<b>Interpretación:</b>", styles['Normal']))
        story.append(Paragraph(interpretation, styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # 3. HALLAZGOS CRÍTICOS
        story.append(Paragraph("2. HALLAZGOS CRÍTICOS", heading_style))
        story.append(Spacer(1, 0.2*inch))
        
        critical_calls = [r for r in batch_result.results if r.critical_findings]
        
        if critical_calls:
            story.append(Paragraph(
                f"Se identificaron <b>{len(critical_calls)} llamadas</b> con hallazgos críticos que requieren atención inmediata:",
                styles['Normal']
            ))
            story.append(Spacer(1, 0.15*inch))
            
            for i, result in enumerate(critical_calls[:10], 1):  # Top 10
                call_info = f"<b>Llamada {i}:</b> {result.audited_call.filename}"
                story.append(Paragraph(call_info, styles['Normal']))
                
                for finding in result.critical_findings[:3]:  # Top 3 findings
                    finding_text = f"  • [{finding.severity.value}] {finding.title}"
                    story.append(Paragraph(finding_text, styles['Normal']))
                
                story.append(Spacer(1, 0.1*inch))
        else:
            story.append(Paragraph(
                "✓ No se detectaron hallazgos críticos en el batch.",
                styles['Normal']
            ))
        
        story.append(Spacer(1, 0.3*inch))
        
        # 4. MÉTRICAS CLAVE
        story.append(Paragraph("3. MÉTRICAS CLAVE", heading_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Distribución de status
        status_counts = {}
        for result in batch_result.results:
            status = result.overall_status
            status_counts[status] = status_counts.get(status, 0) + 1
        
        story.append(Paragraph("<b>Distribución de Status:</b>", styles['Normal']))
        for status, count in status_counts.items():
            pct = (count / batch_result.total_calls) * 100
            story.append(Paragraph(f"  • {status}: {count} ({pct:.1f}%)", styles['Normal']))
        
        story.append(Spacer(1, 0.3*inch))
        
        # 5. RECOMENDACIONES
        story.append(Paragraph("4. RECOMENDACIONES", heading_style))
        story.append(Spacer(1, 0.2*inch))
        
        recommendations = self._get_batch_recommendations(batch_result)
        for i, rec in enumerate(recommendations, 1):
            story.append(Paragraph(f"{i}. {rec}", styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
        
        story.append(Spacer(1, 0.3*inch))
        
        # 6. CONCLUSIÓN OPERATIVA
        story.append(Paragraph("5. CONCLUSIÓN OPERATIVA", heading_style))
        story.append(Spacer(1, 0.2*inch))
        
        conclusion = self._get_batch_conclusion(batch_result)
        story.append(Paragraph(conclusion, styles['Normal']))
        
        # Generar PDF
        doc.build(story)
    
    def _generate_individual_pdf(self, audit_result: AuditResult, output_path: Path):
        """Genera PDF de un audit individual"""
        if not PDF_AVAILABLE:
            raise RuntimeError("reportlab no disponible")
        
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch
        )
        
        story = []
        styles = getSampleStyleSheet()
        
        # Título
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph(self.config.company_name, styles['Title']))
        story.append(Paragraph("REPORTE DE AUDITORÍA INDIVIDUAL", styles['Heading1']))
        story.append(Spacer(1, 0.3*inch))
        
        # Info básica
        call = audit_result.audited_call
        story.append(Paragraph(f"<b>Archivo:</b> {call.filename}", styles['Normal']))
        story.append(Paragraph(f"<b>Call ID:</b> {call.call_id}", styles['Normal']))
        story.append(Paragraph(f"<b>Duración:</b> {call.duration_minutes:.1f} min", styles['Normal']))
        story.append(Paragraph(f"<b>Nivel de Servicio:</b> {call.service_level.value}", styles['Normal']))
        story.append(Paragraph(f"<b>QA Score:</b> {audit_result.qa_score:.1f}%", styles['Normal']))
        story.append(Paragraph(f"<b>Status:</b> {audit_result.overall_status}", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Findings
        if audit_result.findings:
            story.append(Paragraph("HALLAZGOS", styles['Heading2']))
            story.append(Spacer(1, 0.2*inch))
            
            for finding in audit_result.findings:
                severity_color = self._get_severity_color(finding.severity)
                story.append(Paragraph(
                    f"<font color='{severity_color}'>[{finding.severity.value}]</font> <b>{finding.title}</b>",
                    styles['Normal']
                ))
                story.append(Paragraph(finding.description, styles['Normal']))
                if finding.recommendation:
                    story.append(Paragraph(f"<i>Recomendación: {finding.recommendation}</i>", styles['Normal']))
                story.append(Spacer(1, 0.15*inch))
        
        # Métricas
        if audit_result.metrics:
            story.append(Paragraph("MÉTRICAS", styles['Heading2']))
            story.append(Spacer(1, 0.2*inch))
            
            metrics_data = [['Métrica', 'Valor', 'Status']]
            for metric in audit_result.metrics:
                metrics_data.append([
                    metric.name,
                    metric.formatted_value,
                    metric.status.value if metric.status else 'N/A'
                ])
            
            metrics_table = Table(metrics_data, colWidths=[2*inch, 1.5*inch, 1*inch])
            metrics_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(metrics_table)
        
        doc.build(story)
    
    # === DOCX GENERATION ===
    
    def _generate_batch_docx(self, batch_result: BatchAuditResult, output_path: Path):
        """Genera DOCX del batch consolidado"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx no disponible")
        
        doc = Document()
        
        # Título
        title = doc.add_heading(self.config.company_name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('REPORTE CONSOLIDADO DE AUDITORÍA', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Generado: {batch_result.generated_at.strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph(f"Total de llamadas: {batch_result.total_calls}")
        doc.add_page_break()
        
        # 1. Resumen Ejecutivo
        doc.add_heading('1. RESUMEN EJECUTIVO', level=1)
        
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.rows[0].cells[0].text = 'Total de Llamadas'
        table.rows[0].cells[1].text = str(batch_result.total_calls)
        table.rows[1].cells[0].text = 'Llamadas Aprobadas'
        table.rows[1].cells[1].text = str(batch_result.passed_calls)
        table.rows[2].cells[0].text = 'Llamadas Rechazadas'
        table.rows[2].cells[1].text = str(batch_result.failed_calls)
        table.rows[3].cells[0].text = 'Tasa de Aprobación'
        table.rows[3].cells[1].text = f"{batch_result.approval_rate:.1f}%"
        table.rows[4].cells[0].text = 'QA Score Promedio'
        table.rows[4].cells[1].text = f"{batch_result.avg_qa_score:.1f}%"
        table.rows[5].cells[0].text = 'Hallazgos Críticos'
        table.rows[5].cells[1].text = str(batch_result.critical_findings_count)
        table.rows[6].cells[0].text = 'Tiempo de Procesamiento'
        table.rows[6].cells[1].text = f"{batch_result.processing_time_seconds:.1f}s"
        table.rows[7].cells[0].text = 'Llamadas Requieren Atención'
        table.rows[7].cells[1].text = str(len(batch_result.requires_attention))
        
        doc.add_paragraph()
        interpretation = self._get_batch_interpretation(batch_result)
        p = doc.add_paragraph()
        p.add_run('Interpretación: ').bold = True
        p.add_run(interpretation)
        
        # 2. Hallazgos Críticos
        doc.add_heading('2. HALLAZGOS CRÍTICOS', level=1)
        
        critical_calls = [r for r in batch_result.results if r.critical_findings]
        
        if critical_calls:
            doc.add_paragraph(
                f"Se identificaron {len(critical_calls)} llamadas con hallazgos críticos:"
            )
            
            for i, result in enumerate(critical_calls[:10], 1):
                p = doc.add_paragraph(style='List Number')
                p.add_run(f"{result.audited_call.filename}").bold = True
                
                for finding in result.critical_findings[:3]:
                    doc.add_paragraph(
                        f"[{finding.severity.value}] {finding.title}",
                        style='List Bullet 2'
                    )
        else:
            doc.add_paragraph("✓ No se detectaron hallazgos críticos.")
        
        # 3. Métricas Clave
        doc.add_heading('3. MÉTRICAS CLAVE', level=1)
        
        status_counts = {}
        for result in batch_result.results:
            status = result.overall_status
            status_counts[status] = status_counts.get(status, 0) + 1
        
        p = doc.add_paragraph()
        p.add_run('Distribución de Status:').bold = True
        
        for status, count in status_counts.items():
            pct = (count / batch_result.total_calls) * 100
            doc.add_paragraph(f"{status}: {count} ({pct:.1f}%)", style='List Bullet')
        
        # 4. Recomendaciones
        doc.add_heading('4. RECOMENDACIONES', level=1)
        
        recommendations = self._get_batch_recommendations(batch_result)
        for rec in recommendations:
            doc.add_paragraph(rec, style='List Number')
        
        # 5. Conclusión Operativa
        doc.add_heading('5. CONCLUSIÓN OPERATIVA', level=1)
        
        conclusion = self._get_batch_conclusion(batch_result)
        doc.add_paragraph(conclusion)
        
        doc.save(str(output_path))
    
    def _generate_individual_docx(self, audit_result: AuditResult, output_path: Path):
        """Genera DOCX de un audit individual"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx no disponible")
        
        doc = Document()
        
        # Título
        title = doc.add_heading(self.config.company_name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading('REPORTE DE AUDITORÍA INDIVIDUAL', level=1)
        
        # Info básica
        call = audit_result.audited_call
        doc.add_paragraph(f"Archivo: {call.filename}")
        doc.add_paragraph(f"Call ID: {call.call_id}")
        doc.add_paragraph(f"Duración: {call.duration_minutes:.1f} min")
        doc.add_paragraph(f"Nivel: {call.service_level.value}")
        doc.add_paragraph(f"QA Score: {audit_result.qa_score:.1f}%")
        doc.add_paragraph(f"Status: {audit_result.overall_status}")
        
        # Findings
        if audit_result.findings:
            doc.add_heading('HALLAZGOS', level=2)
            
            for finding in audit_result.findings:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"[{finding.severity.value}] {finding.title}").bold = True
                doc.add_paragraph(finding.description)
                if finding.recommendation:
                    doc.add_paragraph(f"→ {finding.recommendation}", style='List Bullet 2')
        
        # Métricas
        if audit_result.metrics:
            doc.add_heading('MÉTRICAS', level=2)
            
            table = doc.add_table(rows=len(audit_result.metrics)+1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            table.rows[0].cells[0].text = 'Métrica'
            table.rows[0].cells[1].text = 'Valor'
            table.rows[0].cells[2].text = 'Status'
            
            for i, metric in enumerate(audit_result.metrics, 1):
                table.rows[i].cells[0].text = metric.name
                table.rows[i].cells[1].text = metric.formatted_value
                table.rows[i].cells[2].text = metric.status.value if metric.status else 'N/A'
        
        doc.save(str(output_path))
    
    # === HELPERS ===
    
    def _get_batch_interpretation(self, batch_result: BatchAuditResult) -> str:
        """Genera interpretación del batch"""
        if batch_result.approval_rate >= 80:
            return (
                f"El rendimiento general es EXCELENTE. {batch_result.approval_rate:.1f}% de las llamadas "
                f"cumplen con los estándares de calidad establecidos. "
                f"El QA Score promedio de {batch_result.avg_qa_score:.1f}% indica un desempeño consistente del equipo."
            )
        elif batch_result.approval_rate >= 60:
            return (
                f"El rendimiento es ACEPTABLE pero requiere atención. {batch_result.approval_rate:.1f}% de aprobación "
                f"indica áreas de mejora. Se recomienda capacitación enfocada en los hallazgos identificados."
            )
        else:
            return (
                f"El rendimiento requiere INTERVENCIÓN INMEDIATA. Solo {batch_result.approval_rate:.1f}% de las llamadas "
                f"aprueban. Se recomienda revisión del protocolo y capacitación urgente del equipo."
            )
    
    def _get_batch_recommendations(self, batch_result: BatchAuditResult) -> List[str]:
        """Genera recomendaciones basadas en resultados"""
        recommendations = []
        
        if batch_result.approval_rate < 70:
            recommendations.append(
                "Implementar programa de capacitación urgente enfocado en protocolos básicos de atención"
            )
        
        if batch_result.critical_findings_count > batch_result.total_calls * 0.3:
            recommendations.append(
                "Establecer sesiones de coaching 1-1 para agentes con hallazgos críticos recurrentes"
            )
        
        if len(batch_result.requires_attention) > batch_result.total_calls * 0.5:
            recommendations.append(
                "Revisar y simplificar scripts de atención para reducir complejidad operativa"
            )
        
        if batch_result.avg_qa_score < 70:
            recommendations.append(
                "Realizar auditoría del proceso de QA para validar criterios de evaluación"
            )
        
        if not recommendations:
            recommendations.append(
                "Mantener el nivel de desempeño actual mediante seguimiento periódico y reconocimiento del equipo"
            )
        
        return recommendations
    
    def _get_batch_conclusion(self, batch_result: BatchAuditResult) -> str:
        """Genera conclusión operativa"""
        if batch_result.total_calls == 0:
            return "No se procesaron llamadas en este batch. Verificar carpeta de entrada."
        
        attention_pct = (len(batch_result.requires_attention) / batch_result.total_calls) * 100
        
        if batch_result.approval_rate >= 80:
            return (
                f"El equipo demuestra un desempeño sobresaliente con {batch_result.approval_rate:.1f}% de aprobación. "
                f"Se recomienda continuar con el monitoreo estándar y reconocer públicamente los logros alcanzados. "
                f"Las {len(batch_result.requires_attention)} llamadas que requieren atención ({attention_pct:.1f}%) "
                f"deben revisarse para identificar oportunidades de mejora continua."
            )
        else:
            return (
                f"Se requiere plan de acción correctivo inmediato. {len(batch_result.requires_attention)} llamadas "
                f"({attention_pct:.1f}%) necesitan seguimiento. Priorizar capacitación en áreas críticas identificadas "
                f"y realizar re-auditoría en 2 semanas para validar mejoras."
            )
    
    def _get_severity_color(self, severity: FindingSeverity) -> str:
        """Obtiene color por severidad"""
        colors_map = {
            FindingSeverity.CRITICAL: '#E74C3C',
            FindingSeverity.HIGH: '#E67E22',
            FindingSeverity.MEDIUM: '#F39C12',
            FindingSeverity.LOW: '#3498DB',
            FindingSeverity.INFO: '#95A5A6'
        }
        return colors_map.get(severity, '#000000')


# Helper functions para uso directo
def generate_batch_reports(
    batch_result: BatchAuditResult,
    output_dir: str = "reports",
    format: str = "both"
) -> Dict[str, str]:
    """
    Helper: Genera reportes consolidados de batch
    
    Args:
        batch_result: Resultado del batch audit
        output_dir: Directorio de salida
        format: 'pdf', 'docx', o 'both'
        
    Returns:
        Dict con paths de reportes generados
        
    Example:
        >>> from daia.application import process_audio_folder
        >>> batch = process_audio_folder("audio_in/")
        >>> reports = generate_batch_reports(batch, format="both")
        >>> print(reports['pdf'])
    """
    config = ReportConfig(output_dir=output_dir)
    generator = ReportGenerator(config)
    return generator.generate_batch_report(batch_result, format)


def generate_individual_reports(
    audit_result: AuditResult,
    output_dir: str = "reports",
    format: str = "both"
) -> Dict[str, str]:
    """
    Helper: Genera reportes de un audit individual
    
    Args:
        audit_result: Resultado individual
        output_dir: Directorio de salida
        format: 'pdf', 'docx', o 'both'
        
    Returns:
        Dict con paths de reportes generados
    """
    config = ReportConfig(output_dir=output_dir)
    generator = ReportGenerator(config)
    return generator.generate_individual_report(audit_result, format)
